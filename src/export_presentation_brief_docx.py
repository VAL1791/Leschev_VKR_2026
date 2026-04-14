from __future__ import annotations

import argparse
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


def add_title(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(18)


def add_bullets(document: Document, lines: list[str]) -> None:
    for line in lines:
        document.add_paragraph(line, style="List Bullet")


def build_core_table(results: pd.DataFrame) -> pd.DataFrame:
    frame = results.copy()
    if "run_status" in frame.columns:
        statuses = frame["run_status"].fillna("done").astype(str).str.strip().str.lower()
        frame = frame[statuses == "done"].copy()
    frame = frame[(frame["topk"] == 10) & (frame["mode"] == "exc")].copy()
    frame = frame[["dataset", "algorithm", "ndcg", "recall", "precision", "map"]]
    return frame.sort_values(["dataset", "ndcg"], ascending=[True, False]).reset_index(drop=True)


def significance_notes(significance: pd.DataFrame) -> list[str]:
    if significance.empty:
        return []
    frame = significance[(significance["metric"] == "ndcg")].copy()
    notes: list[str] = []
    for dataset in sorted(frame["dataset"].dropna().astype(str).unique().tolist()):
        dataset_frame = frame[frame["dataset"] == dataset]
        for ref, comp in [
            ("HybridFeatureRerank", "EASE"),
            ("HybridFeatureRerank", "SequentialTransition"),
        ]:
            row = dataset_frame[
                (dataset_frame["reference_algorithm"] == ref)
                & (dataset_frame["comparison_algorithm"] == comp)
            ]
            if row.empty:
                continue
            item = row.iloc[0]
            significance_label = "значима" if bool(item["significant"]) else "не подтверждена статистически"
            notes.append(
                f"{dataset}: разница {ref} vs {comp} по NDCG@10 {significance_label} "
                f"(delta {item['delta_mean']:+.4f}, p_adj={item['p_value_adj']:.4f})."
            )
    return notes


def build_docx(results: pd.DataFrame, significance: pd.DataFrame, output_path: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    add_title(document, "Краткое представление результата проекта")
    document.add_paragraph(
        "Документ нужен для короткого показа результата: что построено, какие данные использованы, "
        "какие выводы уже можно защищать и как запустить демонстрационный контур."
    )

    document.add_heading("Что сделано", level=1)
    add_bullets(
        document,
        [
            "Собран единый reproducible benchmark для MovieLens 20M, Google Local South Carolina и Goodreads Fantasy 10K.",
            "Построен baseline stack: Popularity, TruncatedSVD, EASE, TFIDF_Content, SequentialTransition и HybridFeatureRerank.",
            "Добавлены robustness-слои: ablation, сегментный анализ, confidence intervals и paired significance checks.",
            "Результаты выводятся в Streamlit dashboard и в defense-ready summary.",
        ],
    )

    document.add_heading("Ключевые результаты", level=1)
    document.add_paragraph("Основной режим сравнения: exc, topk=10, NDCG@10.")
    table_data = build_core_table(results)
    table = document.add_table(rows=1, cols=len(table_data.columns))
    table.style = "Table Grid"
    for cell, name in zip(table.rows[0].cells, table_data.columns):
        cell.text = name
    for row in table_data.itertuples(index=False):
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = f"{value:.4f}" if isinstance(value, float) else str(value)

    document.add_heading("Как интерпретировать", level=1)
    add_bullets(
        document,
        [
            "MovieLens 20M: dense collaborative case; EASE уже очень силён, hybrid статистически не доказывает превосходство.",
            "Google Local South Carolina: strongest applied hybrid case; HybridFeatureRerank значимо превосходит EASE.",
            "Goodreads Fantasy 10K: content-dominant domain; hybrid выигрывает у EASE, но не доказывает превосходство над TFIDF_Content.",
            "SequentialTransition закрывает sequential scope без тяжёлого deep-стека, но не становится главным драйвером качества.",
        ],
    )

    document.add_heading("Что подтверждено статистически", level=1)
    notes = significance_notes(significance)
    if notes:
        add_bullets(document, notes)
    else:
        document.add_paragraph("Статистические проверки не найдены.")

    document.add_heading("Как показать результат", level=1)
    add_bullets(
        document,
        [
            "Открыть dashboard Home и коротко показать project map и текущие takeaways.",
            "На Model Benchmark показать главный leaderboard и confidence/significance tables.",
            "На Robustness & Business показать ablation и объяснить, где важны collaborative, а где content признаки.",
            "На Recommendation Sandbox показать один живой пример истории пользователя и рекомендаций разных моделей.",
        ],
    )

    document.add_heading("Как запускать", level=1)
    add_bullets(
        document,
        [
            "python -m src.run_real_benchmark --refresh",
            "python -m src.validate_dashboard_artifacts --artifacts-dir dashboard/artifacts",
            "python -m src.export_defense_summary",
            "streamlit run dashboard/app.py",
        ],
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main() -> int:
    parser = argparse.ArgumentParser(description="Export a short presentation-oriented DOCX brief.")
    parser.add_argument("--artifacts-dir", default="dashboard/artifacts")
    parser.add_argument("--output", default="docs/RESULT_PRESENTATION_BRIEF.docx")
    args = parser.parse_args()

    project_root = Path(__file__).resolve().parents[1]
    artifacts_dir = Path(args.artifacts_dir)
    output_path = Path(args.output)
    if not artifacts_dir.is_absolute():
        artifacts_dir = (project_root / artifacts_dir).resolve()
    if not output_path.is_absolute():
        output_path = (project_root / output_path).resolve()

    results = pd.read_csv(artifacts_dir / "results.csv")
    significance = pd.read_csv(artifacts_dir / "significance_tests.csv")
    build_docx(results, significance, output_path)
    print(f"[export] presentation brief -> {output_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
